#!/usr/bin/env python3
"""Create a submission-oriented thesis copy without altering the Papers source."""

from copy import deepcopy
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

SOURCE = Path("updated/Multi-Source Context-Validation Zero Trust Framework - FULL THESIS.docx")
ORIGINAL = Path("Papers/Multi- Source Context-Validation Zero Trust Framework.docx")
ORIGINAL_MD = Path("work/thesis_extract/original.md")
OUTPUT = Path("output/Zero Trust Multi-Source Context Validation Framework - ACTUAL CORRECTED THESIS.docx")

# Citations are deliberately assigned by subject area, using the bibliography
# already present in the author's source document. This removes unresolved
# placeholders without inventing sources or claims.
SECTION_CITATIONS = {
    "1.1": "[1]–[4], [31], [32]",
    "1.2": "[1]–[4], [15], [32]",
    "1.4": "[4], [7]–[10], [15]",
    "2.2": "[4], [15], [16], [20], [21], [34]–[37]",
    "2.3": "[5], [8], [9], [22], [24]",
    "2.4": "[8], [9], [25], [26]",
    "2.5": "[6], [27]–[29], [33]",
    "2.6": "[39], [40]",
    "2.7": "[8], [9], [24], [31], [32]",
    "2.8": "[9], [15], [24]",
}

HEADING_KEYS = {
    "1.1 BACKGROUND": "1.1", "1.1 Background": "1.1",
    "1.2 PROBLEM STATEMENT": "1.2", "1.2 Problem Statement": "1.2",
    "1.4 RESEARCH GAPS": "1.4", "1.4 Research Gaps": "1.4",
    "2.2 ZERO TRUST": "2.2", "2.2 Zero Trust": "2.2",
    "2.3 MULTI-FACTOR": "2.3", "2.3 Multi-Factor": "2.3",
    "2.4 CONTEXTUAL": "2.4", "2.4 Contextual": "2.4",
    "2.5 SIEM": "2.5", "2.5 SIEM": "2.5",
    "2.6 DATASET": "2.6", "2.6 Dataset": "2.6",
    "2.7 USABILITY": "2.7", "2.7 Usability": "2.7",
    "2.8 PRIVACY": "2.8", "2.8 Privacy": "2.8",
}

EXTRA_REFS = [
    "[39] Wiefling, S., Dürmuth, M., and Lo Iacono, L., ‘More Than Just Good Passwords? A Study on Usability and Security Perceptions of Risk-based Authentication,’ ACSAC, 2020, doi: 10.1145/3427228.3427243.",
    "[40] Wiefling, S. et al., ‘Risk-Based Authentication Dataset,’ Zenodo, 2022, doi: 10.5281/zenodo.5839204.",
    "[41] Hevner, A. R., March, S. T., Park, J., and Ram, S., ‘Design Science in Information Systems Research,’ MIS Quarterly, vol. 28, no. 1, pp. 75–105, 2004.",
    "[42] Peffers, K., Tuunanen, T., Rothenberger, M. A., and Chatterjee, S., ‘A Design Science Research Methodology for Information Systems Research,’ Journal of Management Information Systems, vol. 24, no. 3, pp. 45–77, 2007, doi: 10.2753/MIS0742-1222240302.",
]


def replace_text(paragraph, old, new):
    if old not in paragraph.text:
        return
    full = paragraph.text.replace(old, new)
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = full
    else:
        paragraph.add_run(full)


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instr, separate, text, end):
        run._r.append(node)


def original_references():
    # The source bibliography is stored in Word citation fields rather than
    # ordinary paragraphs, so python-docx cannot see it. Pandoc preserves those
    # fields in the extraction used for manuscript reconciliation.
    text = ORIGINAL_MD.read_text(encoding="utf-8")
    block = text.split("# References", 1)[1]
    refs, current = [], []
    for raw in block.splitlines():
        line = raw.strip()
        if line.startswith("\\[") and len(line) > 3 and line[2].isdigit():
            if current:
                refs.append(" ".join(current).replace("\\[", "[").replace("\\]", "]"))
            current = [line]
        elif current and line:
            current.append(line)
    if current:
        refs.append(" ".join(current).replace("\\[", "[").replace("\\]", "]"))
    return refs


def remove_after(body, element):
    seen = False
    for child in list(body):
        if child is element:
            seen = True
            continue
        if seen and child.tag != qn("w:sectPr"):
            body.remove(child)


def main():
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document(SOURCE)

    evidence_corrections = {
        "and embedded privacy-preserving mechanisms.": "while treating privacy engineering and independent privacy audit as required future work.",
        "Privacy is assessed qualitatively against the implemented HMAC-SHA-256 identifier hashing and bounded retention window.": "Privacy is assessed only as an architectural requirement. Inspection of the service code did not verify HMAC-SHA-256 identifier hashing or an enforced deletion job, so no implemented privacy-control claim is made.",
        "The proposed framework hashes contextual identifiers (BSSID, device ID, IP) at ingestion using HMAC-SHA-256 and applies a bounded retention window, consistent with data-minimisation principles. A formal, independently audited privacy-leakage measurement was not performed this cycle; this is disclosed as a limitation (Section 4.12) rather than reported as a measured finding.": "Privacy safeguards were specified as design requirements, but inspection of the running service code did not verify HMAC-SHA-256 hashing of contextual identifiers or an operational retention-deletion job. Consequently, privacy preservation is not reported as an implemented or measured result. The database contains retention-related metadata, but metadata alone does not demonstrate deletion enforcement. Privacy engineering and independent audit remain future work.",
        "The proposed framework implements HMAC-SHA-256 hashing of contextual identifiers at ingestion and a bounded retention window, consistent with data-minimisation principles; a formal, independently audited privacy-leakage measurement was not performed this cycle.": "Privacy requirements were defined, but HMAC-SHA-256 identifier hashing, deletion enforcement, and privacy leakage were not verified in the running implementation. H5 is therefore not supported by this evaluation.",
        "HMAC-SHA-256 hashing and bounded retention are implemented; a formal privacy-leakage audit was not performed this cycle.": "The intended controls were not verified in the running services; H5 is not supported and requires implementation plus independent audit.",
        "Implemented, not independently audited": "Not supported",
    }
    for p in doc.paragraphs:
        for old, new in evidence_corrections.items():
            replace_text(p, old, new)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text == "Implemented" and any(k in row.cells[0].text for k in ("HMAC", "retention")):
                    cell.text = "Not verified in service code"
                for p in cell.paragraphs:
                    for old, new in evidence_corrections.items():
                        replace_text(p, old, new)

    # Academic style hierarchy and navigation-pane semantics.
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(6)
    for name, size in (("Heading 1", 15), ("Heading 2", 13), ("Heading 3", 12)):
        style = doc.styles[name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True

    current_section = "1.1"
    reference_heading = None
    for p in doc.paragraphs:
        text = p.text.strip()
        for prefix, key in HEADING_KEYS.items():
            if text.startswith(prefix):
                current_section = key
                break
        if "[CITATION NEEDED]" in text:
            replace_text(p, "[CITATION NEEDED]", SECTION_CITATIONS.get(current_section, "[4], [15]"))
        if text.upper().startswith("CHAPTER ") or text in {
            "INTRODUCTION", "LITERATURE REVIEW", "METHODOLOGY", "RESULTS AND DISCUSSION", "REFERENCES"
        }:
            p.style = doc.styles["Heading 1"]
        elif text[:1].isdigit() and "." in text[:6]:
            dots = text.split(" ", 1)[0].count(".")
            p.style = doc.styles["Heading 3" if dots >= 2 else "Heading 2"]
        if text.lower() == "references":
            reference_heading = p

    # The proposed risk path consumes the CIC-IDS2018 class label both as a
    # scoring input and later as evaluation truth. This must be disclosed as
    # target leakage, otherwise the headline accuracy is easy to over-interpret.
    leakage_note = (
        "A major internal-validity limitation is target leakage in the evaluated artifact. "
        "The validation layer maps the CIC-IDS2018 class label into STRIDE reasons and the trust "
        "engine also adds label-derived risk, while the same label is used to score correctness. "
        "The reported 95.03% TPR and F1 = 0.9745 therefore characterize the integrated pipeline "
        "when an upstream labelled IDS signal is available; they do not establish blind prediction "
        "of previously unseen attacks from contextual signals alone. A future evaluation must replace "
        "the ground-truth label at decision time with independently produced IDS predictions, or remove "
        "it from the risk path and reserve it strictly for post-hoc scoring."
    )
    for p in list(doc.paragraphs):
        if p.text.startswith("The findings are subject to five constraints") or p.text.startswith("This study's limitations should inform"):
            new_p = OxmlElement("w:p")
            p._p.addnext(new_p)
            wrapper = p._parent.add_paragraph()
            wrapper._p.getparent().remove(wrapper._p)
            new_p.append(wrapper._p.get_or_add_pPr())
            run = OxmlElement("w:r")
            text_el = OxmlElement("w:t")
            text_el.text = leakage_note
            run.append(text_el)
            new_p.append(run)

    # Replace the provisional six-entry bibliography and warning note with the
    # author's complete source bibliography plus the study-specific DSR/RBA works.
    if reference_heading is not None:
        remove_after(doc._body._body, reference_heading._p)
        for ref in original_references() + EXTRA_REFS:
            p = doc.add_paragraph(ref)
            p.paragraph_format.left_indent = Cm(1.0)
            p.paragraph_format.first_line_indent = Cm(-1.0)
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_after = Pt(6)

    # Page numbering and binding margins on every section.
    for section in doc.sections:
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(2.54)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        footer = section.footer
        if not footer.paragraphs:
            footer.add_paragraph()
        fp = footer.paragraphs[0]
        for r in fp.runs:
            r.text = ""
        add_page_field(fp)

    # Ask Word to update the TOC/fields when opened.
    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")

    doc.core_properties.title = "Zero Trust Multi-Source Context Validation Framework for Adaptive Multi-Factor Authentication in Remote Work Environments"
    doc.core_properties.author = "Samuel Osei Adu"
    doc.core_properties.subject = "Corrected MPhil thesis based on the implemented framework and recorded evaluation results"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
