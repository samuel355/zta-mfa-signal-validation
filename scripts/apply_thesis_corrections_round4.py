#!/usr/bin/env python3
"""
Round 4 — corrects the Consistency Score (Cs) and Enrichment Trust Score (Es)
description in Section 3.5.3 (paragraphs 330-339) to match what
services/validation/app/enrichment.py and compute_weights() actually implement.

Found during a deep code-vs-thesis audit: the previously-written Es formula
(Ivpn=1 penalty=0.7, Itor=1 penalty=0.9, Imalicious=1 penalty=0.1, Iunknown=1
penalty=0.2) does not correspond to any code anywhere in the codebase — there is
no VPN detection, no Tor exit-node IP check, and no separate malicious/unknown-IP
penalty system. The real enrichment mechanism is a single TLS/JA3 fingerprint
lookup against a curated critical-tag table, contributing one 0.2x weight
discount, not four independently-weighted indicator penalties. Paragraph 331's
"Device-TLS Consistency" binary check also does not exist in the code — Cs is
geographic-only (GPS vs WiFi/IP haversine distance). Paragraph 330's stated
threshold (1000km) also does not match the live deployment (100km, confirmed via
compose/.env DIST_THRESHOLD_KM and the running validation container).
"""
import docx

PATH = "/Users/knight/Apps/multi-source-ztamfa/updated/Multi- Source Context-Validation Zero Trust Framework - CORRECTED.docx"

d = docx.Document(PATH)
paras = d.paragraphs


def set_para_text(idx, new_text):
    p = paras[idx]
    if not p.runs:
        p.add_run(new_text)
        return
    p.runs[0].text = new_text
    for r in p.runs[1:]:
        r.text = ""


set_para_text(330,
    "Where d is the distance between the GPS-reported location and the location "
    "implied by the Wi-Fi access point (preferred) or IP geolocation (fallback, "
    "used only when no Wi-Fi signal is present), and d0 is the threshold "
    "parameter — 100km in the deployed configuration (DIST_THRESHOLD_KM)."
)

set_para_text(331,
    "This is the only consistency check the implemented system performs. An "
    "earlier draft of this section also described a binary Device-TLS "
    "Consistency check (declared device platform vs. TLS fingerprint); no such "
    "check exists anywhere in the codebase, and the claim is removed rather than "
    "retained as an unimplemented aspiration."
)

set_para_text(332,
    "Enrichment Trust Score (Es): In the implemented system this reflects a "
    "single lookup, not a multi-indicator threat-intelligence model: whether the "
    "session's TLS/JA3 fingerprint matches a curated table of known fingerprints "
    "tagged tor_suspect, malware_family_x, scanner_tool, cloud_proxy, "
    "old_openssl, insecure_client, or honeypot_fingerprint (data/tls/"
    "ja3_fingerprints.csv). \"Tor\" detection here is via TLS client "
    "fingerprinting (Tor Browser has a distinctive JA3 signature), not an IP "
    "exit-node list — there is no IP-reputation-based VPN, Tor, malicious-IP, or "
    "unknown-IP check anywhere in the codebase."
)

set_para_text(333,
    "Es = 0.2 if the fingerprint matches a critical tag, else Es = 1.0 (the "
    "corresponding TLS signal's weight is discounted by this factor in the "
    "dynamic weighting step, not zeroed out entirely)."
)

set_para_text(334, "")
set_para_text(335, "")
set_para_text(336, "")
set_para_text(337, "")
set_para_text(338, "")

set_para_text(339,
    "This single penalty value (0.2x) was set heuristically rather than through "
    "formal optimisation — a moderate rather than total discount, reflecting "
    "that a critical TLS tag is suspicious but not, on its own, conclusive proof "
    "of compromise. An earlier draft of this section described four separate "
    "indicator-weighted penalties (Tor 0.9, VPN 0.7, malicious-IP 0.1, "
    "unknown-IP 0.2); that description did not correspond to any implemented "
    "code and has been replaced with the description above. Section 3.5.2 "
    "reports a real sensitivity analysis run against the penalty constants that "
    "do exist in the implementation (missing-signal 0.3x, geographic-mismatch "
    "0.5x, critical-TLS-tag 0.2x)."
)

d.save(PATH)
print("Round 4 Es/Cs correction saved.")
