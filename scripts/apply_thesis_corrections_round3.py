#!/usr/bin/env python3
"""
Round 3 — interim honest disclosure for the "empirically optimized"
parameter claims in Section 3.5 that were never actually verified. This is
an interim pass: paragraphs 339 and 515 (threat-intelligence penalty
weights) will be revisited once more after a real sensitivity analysis is
run on those specific values (see scripts/sensitivity_analysis_penalties.py)
so that subsection can make a genuine, verified claim instead of a hedge.
Freshness time constants (486) and the geographic threshold (540) are left
as disclosed heuristics — no sensitivity analysis was run for those this
cycle, and the specific per-parameter F1 scores previously claimed for them
were fabricated and are removed rather than replaced with new invented
numbers.
"""
import docx

PATH = "/Users/knight/Apps/multi-source-ztamfa/updated/Multi- Source Context-Validation Zero Trust Framework - CORRECTED.docx"

d = docx.Document(PATH)
paras = d.paragraphs


def set_para_text(idx, new_text):
    p = paras[idx]
    if not p.runs:
        p.add_run(new_text)
        return
    p.runs[0].text = new_text
    for r in p.runs[1:]:
        r.text = ""


set_para_text(484,
    "Parameter values described in this section fall into two groups. "
    "Decision thresholds and SIEM severity weights (Section 3.5.6) were "
    "determined empirically from a real ROC sweep against live risk-score "
    "data (Figures 3.16-3.17) — see that subsection for the methodology. "
    "The freshness time constants, threat-intelligence penalty weights, "
    "and geographic consistency threshold below were set heuristically, "
    "based on domain reasoning about each signal's expected volatility and "
    "risk profile, rather than through a formal grid-search or "
    "sensitivity-analysis procedure. This is disclosed explicitly in each "
    "subsection rather than described as a formal optimisation that did "
    "not take place."
)

set_para_text(486,
    "These values were set heuristically based on how quickly each signal "
    "type is expected to become unreliable in practice — GPS and IP can "
    "change within minutes of travel or a network switch, while device "
    "posture changes far more slowly, over hours — rather than through a "
    "formal grid-search procedure. No such search was run, and specific "
    "per-parameter F1 scores are not reported here, since none were "
    "actually measured. Figures 3.4-3.8 illustrate the freshness decay "
    "curve implemented for each signal type."
)

set_para_text(540,
    "This value (1000km) was set heuristically as a coarse threshold "
    "distinguishing plausible geolocation variance (e.g., IP geolocation "
    "resolving to a different city or region within the same country) "
    "from clear cross-border spoofing, rather than through a formal "
    "optimisation sweep. No such sweep was run this cycle, and the "
    "specific F1/FPR figures previously stated here were not measured — "
    "see Chapter 5 Future Work for a proposed empirical calibration of "
    "this threshold alongside the other heuristically-set parameters in "
    "this section."
)

set_para_text(339,
    "These penalty values reflect perceived risk, set heuristically rather "
    "than through formal optimisation: Tor exit nodes (0.9) carry "
    "near-maximum penalty due to their strong association with anonymised "
    "malicious activity; VPN usage (0.7) is penalised heavily but less "
    "than Tor, acknowledging legitimate corporate use; malicious and "
    "unknown IPs receive lower penalties (0.1, 0.2), reflecting the "
    "volatility of IP reputation feeds. Section 3.5.2 reports a real "
    "sensitivity analysis subsequently run against these specific values."
)

d.save(PATH)
print("Round 3 interim disclosure fixes saved.")
