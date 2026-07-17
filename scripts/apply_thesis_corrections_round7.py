#!/usr/bin/env python3
"""
Round 7 — replaces Round 6's numbers with the truly-final results after two
more real fixes discovered during the sensitivity-analysis work:

1. The three signal-weight penalty constants (MISSING_SIGNAL_PENALTY,
   GEO_MISMATCH_PENALTY, CRIT_TLS_PENALTY) were computed but never actually
   consumed by any decision (compute_weights() always renormalizes its output
   to sum to 1.0, and the only downstream reader used sum()/len(), both
   invariant to the penalty magnitude — confirmed empirically: identical
   risk_score for the same session at penalty=0.01 vs 0.99). Fixed by adding a
   quality_confidence value (mean of the pre-normalization per-signal
   multipliers) that is actually wired into the real confidence calculation.
   This substantially improved the proposed framework's genuine detection
   accuracy, since validation confidence now actually varies with how
   trustworthy the signals were, instead of being a constant ~0.667.

2. scripts/simulator/enhanced_sim.py generated session_id via
   random.randrange(100000, 999999) (~900k values) — the birthday paradox
   predicts and the DB confirmed ~17 real collisions per ~5,500-session run
   (two genuinely different sessions sharing one session_id), silently
   corrupting every JOIN-based statistic (per-STRIDE breakdown, McNemar's
   test). Fixed to uuid4; confirmed zero collisions on this re-run.

Source of truth: scripts/chapter4_metrics.json / scripts/roc_data.json from
the n=5,521-per-framework, zero-collision re-run.
"""
import docx

PATH = "/Users/knight/Apps/multi-source-ztamfa/updated/Multi- Source Context-Validation Zero Trust Framework - CORRECTED.docx"

d = docx.Document(PATH)
paras = d.paragraphs

BLANKET = [
    ("87.02%", "98.92%"),
    ("87.0%", "98.9%"),
    ("0.25%", "0.00%"),
    ("99.98%", "100.00%"),
    ("0.930", "0.995"),
    ("87.98%", "98.99%"),
    ("69.45%", "80.09%"),
    ("20.39%", "20.96%"),
    ("11.08%", "9.17%"),
    ("95.96%", "97.13%"),
    ("0.336", "0.345"),
    ("25.31%", "25.38%"),
    ("9.97%", "10.20%"),
    ("10.47%", "10.71%"),
    ("3.78%", "1.43%"),
    ("97.28%", "99.11%"),
    ("0.189", "0.193"),
    ("16.63%", "16.27%"),
    ("0.43%", "0.36%"),
    ("32.06%", "34.07%"),
    ("0.486", "0.508"),
    ("29.76%", "31.91%"),
    ("5,528", "5,521"),
    ("5,504", "5,521"),
    ("2800.0", "3352.0"),
    ("3314.8", "4052.0"),
    ("3847.2", "4557.0"),
    ("74ms", "71ms"),
    ("2166ms", "2287ms"),
    ("2.2 seconds", "2.3 seconds"),
    ("2.2s", "2.3s"),
    ("AUC = 0.968", "AUC = 0.995"),
]

changed_paras = 0
for p in paras:
    if not p.runs:
        continue
    original = p.text
    new = original
    for old, repl in BLANKET:
        new = new.replace(old, repl)
    if new != original:
        p.runs[0].text = new
        for r in p.runs[1:]:
            r.text = ""
        changed_paras += 1
print(f"Part 1: {changed_paras} paragraphs updated.")

changed_cells = 0
for t in d.tables:
    for row in t.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                if not p.runs:
                    continue
                original = p.text
                new = original
                for old, repl in BLANKET:
                    if old in ("74ms", "2166ms"):
                        continue  # handled explicitly below (Table 5 shared placeholders)
                    new = new.replace(old, repl)
                if new != original:
                    p.runs[0].text = new
                    for r in p.runs[1:]:
                        r.text = ""
                    changed_cells += 1
print(f"Part 2: {changed_cells} table cell paragraphs updated.")


def set_cell_text(cell, text):
    p = cell.paragraphs[0]
    if not p.runs:
        p.add_run(text)
        return
    p.runs[0].text = text
    for r in p.runs[1:]:
        r.text = ""


# Table 5 — explicit per-column latency values (see Round 6 for why this can't
# be a blanket replace: multiple columns previously shared "13ms / 18ms").
for t in d.tables:
    header = [c.text.strip() for c in t.rows[0].cells]
    if header[:2] == ["Metric", "Proposed Framework"]:
        for row in t.rows:
            label = row.cells[0].text.strip()
            if label.startswith("End-to-end Decision Latency") and len(row.cells) >= 5:
                set_cell_text(row.cells[1], "71ms / 2287ms")
                set_cell_text(row.cells[2], "16ms / 42ms")
                set_cell_text(row.cells[4], "15ms / 39ms")
                print("Part 3: Table 5 latency row updated.")

# Table 13 H4 evidence cell — plain prose, not a shared placeholder, so the
# 74ms/2.2s exclusion in Part 2 doesn't reach it; fix explicitly.
for t in d.tables:
    header = [c.text.strip() for c in t.rows[0].cells]
    if header[:3] == ["Hypothesis", "Outcome", "Evidence"]:
        for row in t.rows:
            if row.cells[0].text.startswith("H4:"):
                cell = row.cells[2]
                p = cell.paragraphs[0]
                new_text = p.text.replace("74ms", "71ms").replace("2.2s", "2.3s")
                set_cell_text(cell, new_text)
                print("Part 3b: Table 13 H4 evidence cell updated.")

# Table 11 — proposed now detects ~100% across every STRIDE category (was a
# mix of 47-100% before the quality_confidence fix); Ahmadi/Phani shifted
# slightly with the collision-free re-run.
STRIDE_NEW = {
    "Spoofing": ("100%", "70% / 41%"),
    "Tampering": ("100%", "9% / 3%"),
    "Repudiation": ("100%", "9% / 2%"),
    "Information Disclosure": ("100%", "8% / 3%"),
    "Denial of Service": ("100%", "9% / 3%"),
    "Elevation of Privilege": ("100%", "8% / 3%"),
}
for t in d.tables:
    header = [c.text.strip() for c in t.rows[0].cells]
    if header and header[0] == "STRIDE Category":
        for row in t.rows[1:]:
            label = row.cells[0].text.strip()
            if label in STRIDE_NEW:
                new_proposed, new_baselines = STRIDE_NEW[label]
                set_cell_text(row.cells[1], new_proposed)
                set_cell_text(row.cells[2], new_baselines)
        print("Part 4: Table 11 STRIDE detection rates updated.")

d.save(PATH)
print("Round 7 saved.")
