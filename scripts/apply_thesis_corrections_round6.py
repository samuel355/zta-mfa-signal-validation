#!/usr/bin/env python3
"""
Round 6 — replaces every Chapter 4/5 headline number with the results of the
clean re-run performed after removing all random-number contamination from
services/trust/app/decision_engine.py, services/ablation/app/ablation_engine.py,
services/ahmadi2025/app/main.py, and services/phani2025/app/main.py (see git
history for those fixes). Source of truth: scripts/chapter4_metrics.json and
scripts/roc_data.json from the n=5,528-per-framework re-run.

Old -> new headline numbers:
  proposed:  TPR 88.30%->87.02%, FPR 2.86%->0.25%, Precision 99.83%->99.98%,
             F1 0.937->0.930, detection-accuracy 88.75%->87.98%,
             step-up 70.79%->69.45%, median latency 47ms->74ms,
             p95 latency 2142ms->2166ms (~2.1s->~2.2s)
  ablation:  TPR 52.13%->32.06%, FPR 20.95%->0.00%, F1 0.680->0.486,
             detection-accuracy n/a, step-up 42.60%->29.76%,
             median/p95 latency 13/18ms->16/39ms
  ahmadi:    TPR 21.86%->20.39%, FPR 5.71%->11.08%, Precision 98.61%->95.96%,
             F1 0.358->0.336, detection-accuracy 25.56%->25.31%,
             step-up 10.91%->9.97%, median/p95 latency 13/18ms->14/36ms,
             Spoofing STRIDE detection 72%->70%
  phani:     TPR 6.11%->10.47%, FPR 1.90%->3.78%, Precision 98.35%->97.28%,
             F1 0.115->0.189, detection-accuracy 10.81%->16.63%,
             step-up 5.06%->0.43%, median/p95 latency 13/18ms->14/35ms,
             Spoofing STRIDE detection 24%->41%
  n = 2,054 -> 5,528 per framework; McNemar paired n = 2,052 -> 5,504
  McNemar chi2: 557.5->2800.0 (ablation), 1261.9->3314.8 (ahmadi),
                1585.1->3847.2 (phani); all still p<0.001
"""
import docx

PATH = "/Users/knight/Apps/multi-source-ztamfa/updated/Multi- Source Context-Validation Zero Trust Framework - CORRECTED.docx"

d = docx.Document(PATH)
paras = d.paragraphs

# ---------------------------------------------------------------------------
# Part 1: safe, unique, document-wide substring replacements (paragraphs only
# — table cells handled explicitly in Part 3 to avoid cross-column collisions
# like Table 5's three identical "13ms / 18ms" placeholder cells).
# ---------------------------------------------------------------------------
BLANKET = [
    ("88.30%", "87.02%"),
    ("88.3%", "87.0%"),
    ("21.86%", "20.39%"),
    ("6.11%", "10.47%"),
    ("2.86%", "0.25%"),
    ("5.71%", "11.08%"),
    ("1.90%", "3.78%"),
    ("99.83%", "99.98%"),
    ("98.61%", "95.96%"),
    ("98.35%", "97.28%"),
    ("0.937", "0.930"),
    ("0.358", "0.336"),
    ("0.115", "0.189"),
    ("88.75%", "87.98%"),
    ("25.56%", "25.31%"),
    ("10.81%", "16.63%"),
    ("52.13%", "32.06%"),
    ("20.95%", "0.00%"),
    ("0.680", "0.486"),
    ("42.60%", "29.76%"),
    ("70.79%", "69.45%"),
    ("10.91%", "9.97%"),
    ("5.06%", "0.43%"),
    ("2,054", "5,528"),
    ("2,052", "5,504"),
    ("557.5", "2800.0"),
    ("1261.9", "3314.8"),
    ("1585.1", "3847.2"),
    ("47ms", "74ms"),
    ("2142ms", "2166ms"),
]

changed_paras = 0
for p in paras:
    if not p.runs:
        continue
    original = p.text
    new = original
    for old, repl in BLANKET:
        new = new.replace(old, repl)
    # narrow, paragraph-scoped fixes that would collide with unrelated
    # literature-review content (Table 0) if done as blanket replacements
    new = new.replace("2.1 seconds", "2.2 seconds").replace("2.1s", "2.2s")
    new = new.replace("72% and 24% respectively", "70% and 41% respectively")
    if new != original:
        p.runs[0].text = new
        for r in p.runs[1:]:
            r.text = ""
        changed_paras += 1

print(f"Part 1: {changed_paras} paragraphs updated by blanket substring match.")

# ---------------------------------------------------------------------------
# Part 2: table cells — blanket-safe values only (no ambiguous latency/percent
# strings that collide across columns).
# ---------------------------------------------------------------------------
TABLE_BLANKET = [x for x in BLANKET if x[0] not in ("47ms", "2142ms")]

changed_cells = 0
for t in d.tables:
    for row in t.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                if not p.runs:
                    continue
                original = p.text
                new = original
                for old, repl in TABLE_BLANKET:
                    new = new.replace(old, repl)
                if new != original:
                    p.runs[0].text = new
                    for r in p.runs[1:]:
                        r.text = ""
                    changed_cells += 1

print(f"Part 2: {changed_cells} table cell paragraphs updated by blanket match.")

# ---------------------------------------------------------------------------
# Part 3: Table 5 (End-to-end latency). Row layout confirmed by direct
# inspection: [label, proposed, ablation, excluded(Jimmy), baseline]. The
# "baseline" column already collapses Ahmadi/Phani into one value elsewhere in
# this row (e.g. "Architecture" row uses identical text for both) — their new
# latencies are close enough (14ms/36ms vs 14ms/35ms) to keep that convention
# rather than restructuring the table.
# ---------------------------------------------------------------------------
def set_cell_text(cell, text):
    p = cell.paragraphs[0]
    if not p.runs:
        p.add_run(text)
        return
    p.runs[0].text = text
    for r in p.runs[1:]:
        r.text = ""

for t in d.tables:
    header = [c.text.strip() for c in t.rows[0].cells]
    if header[:2] == ["Metric", "Proposed Framework"]:
        for row in t.rows:
            label = row.cells[0].text.strip()
            if label.startswith("End-to-end Decision Latency") and len(row.cells) >= 5:
                set_cell_text(row.cells[1], "74ms / 2166ms")
                set_cell_text(row.cells[2], "16ms / 39ms")
                set_cell_text(row.cells[4], "14ms / 36ms")
                print("Part 3: Table 5 latency row updated.")

# ---------------------------------------------------------------------------
# Part 4: Table 11 (per-STRIDE detection: Proposed vs Ahmadi/Phani)
# ---------------------------------------------------------------------------
STRIDE_NEW = {
    "Spoofing": ("100%", "70% / 41%"),
    "Tampering": ("47%", "11% / 4%"),
    "Repudiation": ("49%", "6% / 3%"),
    "Information Disclosure": ("100%", "10% / 4%"),
    "Denial of Service": ("99%", "8% / 3%"),
    "Elevation of Privilege": ("100%", "9% / 2%"),
}
for t in d.tables:
    header = [c.text.strip() for c in t.rows[0].cells]
    if header and header[0] == "STRIDE Category":
        for row in t.rows[1:]:
            label = row.cells[0].text.strip()
            if label in STRIDE_NEW:
                new_proposed, new_baselines = STRIDE_NEW[label]
                set_cell_text(row.cells[1], new_proposed)
                set_cell_text(row.cells[2], new_baselines)
        print("Part 4: Table 11 STRIDE detection rates updated.")

d.save(PATH)
print("Round 6 saved.")
