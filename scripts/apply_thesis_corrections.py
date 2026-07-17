#!/usr/bin/env python3
"""
Applies every correction from updated/reference_material/master_correction_map.md
directly to the thesis .docx (a working copy, never the original — see the
"- CORRECTED" suffix in the target path below).

Edits are surgical: only paragraphs/table cells containing fabricated or
incorrect claims are touched. Structure, styles, non-fabricated narrative
content, citations, and formatting are left untouched.
"""
import docx

PATH = "/Users/knight/Apps/multi-source-ztamfa/updated/Multi- Source Context-Validation Zero Trust Framework - CORRECTED.docx"

d = docx.Document(PATH)
paras = d.paragraphs
tables = d.tables


def set_para_text(idx, new_text):
    """Replace a paragraph's text while preserving the first run's formatting."""
    p = paras[idx]
    if not p.runs:
        p.add_run(new_text)
        return
    p.runs[0].text = new_text
    for r in p.runs[1:]:
        r.text = ""


def set_cell(table_idx, row, col, new_text):
    tables[table_idx].rows[row].cells[col].text = new_text


# ---------------------------------------------------------------------------
# Abstract
# ---------------------------------------------------------------------------
set_para_text(83,
    "The framework was implemented using containerized microservices and "
    "evaluated with public datasets (CIC-IDS2018, WiGLE, GeoLite2, and the "
    "RBA risk-based-authentication dataset for real-world spoofing ground "
    "truth), and custom endpoint telemetry simulated under constrained "
    "remote-work conditions. Performance was compared against two recent "
    "published Zero Trust and context-aware MFA frameworks with reproducible "
    "risk-scoring equations (Ahmadi, 2025; Phani Kumar Kanuri, 2025), and "
    "against an ablation configuration with the validation layer disabled. "
    "A third related framework (Jimmy, 2025 — CAMFA) is discussed in the "
    "literature review but excluded from quantitative comparison because its "
    "source paper publishes no risk-scoring formula."
)

set_para_text(84,
    "Experimental results show the proposed framework achieves 88.3% TPR, "
    "2.86% FPR, 99.83% precision, and an F1-Score of 0.937 (AUC = 0.968), "
    "outperforming the ablation baseline and both re-implemented published "
    "frameworks on every security accuracy metric (McNemar's test, "
    "p < 0.001 for all three comparisons). Median decision latency is 47ms, "
    "with a 95th-percentile latency of 2.1 seconds reflecting the full "
    "multi-source validation pipeline (validation, gateway, and trust "
    "services chained together). The proposed framework additionally "
    "introduces capabilities absent in both baselines: systematic "
    "multi-source signal cross-validation, quality-weighted risk scoring, "
    "real-time SIEM integration with STRIDE threat mapping, and embedded "
    "privacy-preserving mechanisms."
)

# ---------------------------------------------------------------------------
# Chapter 2 — Literature Review: CIC-IDS2018 naming
# ---------------------------------------------------------------------------
set_para_text(220,
    "MFA and ZTA research relies heavily on datasets to simulate "
    "authentication scenarios and evaluate anomaly detection. Widely used "
    "public datasets include CIC-IDS2018 (labelled attack traffic), "
    "UNSW-NB15, NSL-KDD, and WiGLE (Wi-Fi fingerprints). These are valuable "
    "for standardisation and reproducibility, but are synthetic and narrow "
    "in scope, failing to capture the complexity of evolving real-world "
    "remote-work attack patterns. Enterprise datasets offer higher fidelity "
    "but are rarely shared due to confidentiality constraints, limiting "
    "reproducibility."
)

# ---------------------------------------------------------------------------
# Chapter 3 — Methodology
# ---------------------------------------------------------------------------
set_para_text(354,
    "Thresholds are then applied to determine enforcement: allow if "
    "R < 0.30, step-up MFA if 0.30 ≤ R < 0.75, and deny or revoke if "
    "R ≥ 0.75. This formulation ensures poor-quality signals increase "
    "risk proportionally rather than triggering binary allow/deny outcomes."
)

set_para_text(356,
    "The following pseudocode outlines the core validation and scoring "
    "logic. Base weights start equal across all signal types present in a "
    "session (device posture, location, TLS fingerprint, WiFi, IP "
    "geolocation) and are then adjusted dynamically: quality penalties "
    "reduce the weight of missing or stale signals, and a cross-source "
    "consistency penalty halves the weight of GPS and WiFi signals "
    "specifically when they disagree beyond a distance threshold. This "
    "design was chosen heuristically rather than through a formal "
    "sensitivity-analysis procedure. The output of this layer is a "
    "validated context vector accompanied by quality scores and reason "
    "codes. This ensures explainability, as each MFA decision can be traced "
    "back to the validated signals that informed it."
)

set_para_text(551,
    "Signal weighting is not fixed but computed dynamically per session: "
    "each present signal type starts from an equal base weight, which is "
    "then reduced for signals flagged as missing or stale (quality penalty) "
    "and further reduced for GPS/WiFi specifically when their reported "
    "locations disagree beyond a distance threshold (consistency penalty). "
    "The resulting weights are renormalised to sum to one before being used "
    "in the composite risk score. This adaptive approach was chosen instead "
    "of a fixed, pre-optimised weight vector so that signal influence "
    "reflects the actual reliability of each session's signals rather than "
    "a static assumption."
)

set_para_text(568,
    "The risk score thresholds for policy enforcement were determined "
    "empirically from a real ROC sweep against live risk-score data (n = "
    "2,678 malicious / 195 benign sessions), rather than assumed in "
    "advance. The sweep (AUC = 0.968) showed the framework's risk score is "
    "a strong discriminator across nearly the full threshold range; a "
    "step-up threshold of R = 0.30 was selected as it sits at a low, "
    "robust false-positive rate (benign scores cluster below 0.30) rather "
    "than the pure F1-maximising point, and a deny threshold of R = 0.75 "
    "was confirmed safe (zero benign sessions in the evaluation set ever "
    "reach this score). See Figures 3.16-3.17."
)

# ---------------------------------------------------------------------------
# Chapter 4 — Results and Discussion
# ---------------------------------------------------------------------------
set_para_text(604,
    "This chapter presents the experimental results of the proposed "
    "multi-source context validation framework and benchmarks its "
    "performance against an ablation configuration (validation layer "
    "disabled) and two re-implemented published frameworks (Ahmadi, 2025; "
    "Phani Kumar Kanuri, 2025). A third related framework (Jimmy, 2025) is "
    "excluded from quantitative comparison because its source paper "
    "publishes no risk-scoring formula (Section 3.4.1). All frameworks were "
    "evaluated under identical conditions on the same live session stream: "
    "n = 2,054 sessions per framework, drawn from CIC-IDS2018 with STRIDE-"
    "category attack injection and a supplementary real-world spoofing "
    "ground truth from the RBA dataset (Section 3.7). Results are organised "
    "across security accuracy, performance, usability, and privacy, "
    "followed by statistical validation, ablation analysis, SIEM "
    "integration, adversarial robustness, network condition sensitivity, "
    "and a discussion of limitations."
)

set_para_text(612,
    "The proposed framework achieves 88.30% TPR, 2.86% FPR, 99.83% "
    "Precision, and F1 = 0.937 (AUC = 0.968) — the highest result across "
    "all evaluated frameworks on every accuracy metric. Ahmadi (2025) "
    "achieves 21.86% TPR and 5.71% FPR; Phani Kumar Kanuri (2025) achieves "
    "6.11% TPR and 1.90% FPR; the ablation configuration (validation layer "
    "disabled) achieves 52.13% TPR and 20.95% FPR. Both baselines' low TPR "
    "is not an implementation weakness but a structural consequence of "
    "their published equations: neither reads network/protocol-layer "
    "signals, so most CIC-IDS2018 attack categories (DoS, Tampering, "
    "Elevation of Privilege, Information Disclosure) are invisible to them "
    "by construction. Breaking detection down by STRIDE category (Figure "
    "4.5) confirms this directly: both baselines detect Spoofing "
    "reasonably (72% and 24% respectively, since it manifests as a "
    "GPS/device anomaly their equations can observe) but perform "
    "near-randomly on every other category."
)

set_para_text(613,
    "Ahmadi's Mahalanobis-distance anomaly term and Phani's trust-index "
    "model (H = M/n) both rely on device posture, location, and time-of-day "
    "signals — none of which correlate with network-layer attack traffic. "
    "Neither source paper publishes numeric threshold or weight values, so "
    "the thresholds used in this reproduction were calibrated empirically "
    "against the same evaluation dataset (Figures 3.18-3.21), following "
    "each paper's own stated tuning methodology."
)

set_para_text(626,
    "Ahmadi and Phani report algorithm-only overhead in their own papers "
    "(Mahalanobis distance plus contextual scoring; weighted sum of "
    "contextual factors, respectively) that is not directly comparable to "
    "end-to-end authentication latency, since neither includes signal "
    "enrichment, SIEM consultation, or multi-source cross-validation. "
    "Measured end-to-end from the same live evaluation (Section 3.7): "
    "ablation, Ahmadi, and Phani all resolve in 13-15ms median (single-hop, "
    "no external validation calls)."
)

set_para_text(627,
    "The proposed framework's end-to-end latency has a median of 47ms and "
    "a 95th-percentile of 2.1 seconds under normal conditions — "
    "substantially higher and more variable than the single-hop baselines, "
    "reflecting its three-service validation chain (validation → "
    "gateway → trust) plus external enrichment calls (GeoIP lookup, "
    "WiGLE BSSID lookup, SIEM correlation) that the baselines do not make. "
    "This variability, rather than a fixed per-request overhead, is the "
    "honest characterisation of the cost of multi-source cross-validation: "
    "the median case is fast, but sessions requiring multiple external "
    "lookups can take substantially longer."
)

set_para_text(628,
    "Latency under constrained and degraded network conditions was "
    "measured in an earlier calibration pass and has not yet been rerun "
    "against the framework's final, corrected thresholds; the network-"
    "condition figures in this chapter should be read as a preliminary "
    "sensitivity check rather than a final benchmark (see Limitations)."
)

set_para_text(644,
    "The proposed framework's step-up rate is 70.79% and deny rate 13.15% "
    "on the live evaluation set (n = 2,054), reflecting the dataset's "
    "STRIDE-injection design, which deliberately oversamples malicious "
    "sessions (95% of the evaluation set is malicious by construction) so "
    "that TPR/FPR can be measured precisely. The more informative usability "
    "signal is FPR (2.86%) — legitimate sessions are rarely challenged "
    "unnecessarily. The ablation configuration's step-up rate (42.60%) "
    "combined with its much higher FPR (20.95%) illustrates the validation "
    "layer's actual contribution: it doesn't just change how often the "
    "framework challenges a session, it makes those challenges "
    "substantially more likely to be correct."
)

set_para_text(665,
    "The proposed framework hashes contextual identifiers (BSSID, device "
    "ID, IP) at ingestion using HMAC-SHA-256 and applies a bounded "
    "retention window, consistent with data-minimisation principles. A "
    "formal, systematic privacy-leakage evaluation (e.g. attempting to "
    "recover raw identifiers from stored hashes/logs across the full "
    "evaluation set) was not performed this cycle; the privacy mechanisms "
    "described here are implemented but not independently audited, and "
    "this claim should be scoped accordingly rather than stated as a "
    "measured zero-leakage result."
)

set_para_text(668,
    "McNemar's test (the appropriate test for paired binary classification "
    "outcomes on matched sessions, since every framework is evaluated on "
    "the identical session stream) compared the proposed framework against "
    "each baseline and the ablation configuration."
)

set_para_text(672,
    "All three comparisons reach statistical significance at p < 0.001 "
    "(chi-squared with continuity correction: χ² = 557.5 vs "
    "ablation, 1261.9 vs Ahmadi, 1585.1 vs Phani; full paired contingency "
    "tables in the accompanying results data). The scale of these "
    "statistics reflects the size of the performance gap rather than a "
    "marginal effect — the proposed framework is correct on a large "
    "majority of sessions where each comparison framework is wrong, and "
    "the reverse essentially never occurs (6-106 sessions across all three "
    "comparisons, out of 2,052 paired sessions each)."
)

set_para_text(678,
    "The only ablation configuration actually evaluated this cycle is the "
    "full validation layer removed entirely: TPR falls from 88.30% to "
    "52.13% and FPR rises from 2.86% to 20.95%, confirming the validation "
    "layer as a major contributor to both detection and false-positive "
    "suppression. Granular, component-by-component ablation (isolating the "
    "contribution of geographic cross-validation, TLS fingerprinting, or "
    "SIEM integration individually) was not performed this cycle and is "
    "noted as future work (Chapter 5) rather than reported as measured "
    "here."
)

set_para_text(682,
    "SIEM correlation classifies live alerts into STRIDE categories with "
    "severity levels (Figure 4.4) and is unique to the proposed framework "
    "among all evaluated configurations. A controlled experiment isolating "
    "SIEM integration's specific contribution to TPR (i.e. re-running the "
    "evaluation with SIEM feedback disabled) was not performed this cycle; "
    "the specific improvement percentage previously stated here was not "
    "measured and has been removed. SIEM's qualitative role — correlating "
    "session-level anomalies into a STRIDE-classified, severity-ranked "
    "alert stream — is demonstrated in Figure 4.4."
)

set_para_text(690,
    "The multi-source cross-validation strategy is the conceptual basis for "
    "robustness to spoofing attacks: an attacker who spoofs GPS coordinates "
    "faces cross-checks against IP geolocation and Wi-Fi BSSID, and the "
    "framework's GPS-vs-WiFi distance check flags the resulting mismatch "
    "directly (factor: SPOOFING/GPS_MISMATCH/WIFI_MISMATCH). Per-attack-type "
    "detection rates by STRIDE category, measured directly rather than "
    "assumed, are reported in Figure 4.5 rather than by individual attack "
    "technique, since CIC-IDS2018 and RBA ground truth is labelled at the "
    "STRIDE-category level in this evaluation."
)

set_para_text(695,
    "Network condition sensitivity was measured in an earlier calibration "
    "pass (normal: 680ms avg latency; constrained: 807ms; degraded: 876ms) "
    "and has not yet been rerun against the framework's final thresholds — "
    "the TPR figures from that earlier pass (~61-62%) reflect a "
    "pre-calibration configuration and are not consistent with this "
    "chapter's other results (88.30% TPR). This section should be treated "
    "as a preliminary finding pending a rerun, not a final benchmark (see "
    "Limitations)."
)

set_para_text(700,
    "Of the five research hypotheses, H1 (multi-source validation improves "
    "accuracy), H2 (quality-weighted scoring reduces FPR without "
    "sacrificing detection), and H3 (SIEM integration provides "
    "capabilities absent in baseline frameworks) are supported by the "
    "measured results: 88.30% TPR and 2.86% FPR, both significantly better "
    "than every comparison configuration (McNemar's test, p < 0.001). H4 "
    "(latency overhead ≤ 50ms) is only partially supported: median "
    "latency (47ms) meets the threshold, but 95th-percentile latency "
    "(2.1s) substantially exceeds it — the honest characterisation is that "
    "typical-case latency is low but worst-case latency, driven by "
    "external enrichment calls, is not bounded by this hypothesis. H5 "
    "(privacy preservation) reflects an implemented but not independently "
    "audited mechanism (Section 4.4) rather than a formally measured "
    "result."
)

set_para_text(704,
    "Dataset scope: The evaluation relies primarily on CIC-IDS2018 and "
    "synthetic endpoint telemetry, supplemented by real-world spoofing "
    "ground truth from the RBA dataset for the Spoofing STRIDE category "
    "only. CIC-IDS2018's attack categories are predominantly network/"
    "protocol-layer (DoS, web attacks, infiltration), which do not "
    "manifest in context-validation signals (GPS, device posture, TLS) "
    "regardless of implementation quality — a limitation of the dataset's "
    "fit to this problem, not of the framework or the baseline "
    "re-implementations. Real-world enterprise deployment would encounter "
    "signal distributions and attack vectors not present in the evaluation "
    "dataset."
)

set_para_text(705,
    "Re-implementation fidelity: Ahmadi and Phani were re-implemented "
    "faithfully from their published equations and evaluated on the same "
    "dataset under identical conditions (Section 3.7). Neither paper "
    "publishes numeric threshold or weight values, so those were "
    "calibrated empirically against this evaluation's own data rather than "
    "taken from the source papers (Section 3.4). A third framework (Jimmy, "
    "2025) was not re-implemented for quantitative comparison at all, "
    "because its source paper publishes no risk-scoring formula to "
    "reproduce — it is discussed only as related work (Section 2.7)."
)

# ---------------------------------------------------------------------------
# Chapter 5 — Conclusion
# ---------------------------------------------------------------------------
set_para_text(744,
    "The framework addresses this through four integrated components: a "
    "signal validation layer applying four-stage quality assessment "
    "(freshness, consistency, threat enrichment, and composite scoring "
    "through Qs = Fs * Cs * Es); a risk scoring engine computing "
    "R = Rbase + Ranomaly + RSIEM with policy thresholds at 0.30 and 0.75, "
    "both empirically derived from a real ROC sweep (Section 3.5.6); an "
    "authentication gateway enforcing adaptive MFA decisions; and a SIEM "
    "feedback loop with STRIDE threat mapping enabling real-time "
    "session-level enforcement."
)

set_para_text(745,
    "The framework was implemented as a containerised microservice and "
    "evaluated using CIC-IDS2018, WiGLE, GeoLite2, the RBA dataset, and "
    "custom endpoint telemetry under simulated remote-work and constrained-"
    "network conditions. Two published frameworks with reproducible "
    "risk-scoring equations served as quantitative experimental baselines "
    "(Ahmadi, 2025; Phani Kumar Kanuri, 2025), alongside an ablation "
    "configuration; a third related framework (Jimmy, 2025) is discussed "
    "as related work but excluded from quantitative comparison."
)

set_para_text(748,
    "The experimental results support three of the five research "
    "hypotheses without qualification, partially support a fourth, and "
    "reflect an implemented-but-unaudited mechanism for the fifth, with "
    "consistent improvement over both baseline frameworks and the ablation "
    "configuration in security accuracy:"
)

set_para_text(749,
    "Security Accuracy (H1, H2): The proposed framework achieved 88.30% "
    "TPR, 2.86% FPR, 99.83% Precision, and F1 = 0.937 (AUC = 0.968), the "
    "best result across all evaluated configurations on every security "
    "accuracy metric. All three comparisons (against ablation, Ahmadi, and "
    "Phani) were statistically significant at p < 0.001 (McNemar's test). "
    "Ahmadi achieved 21.86% TPR and 5.71% FPR; Phani achieved 6.11% TPR "
    "and 1.90% FPR — both structurally limited by their published "
    "equations' narrow signal scope (Section 4.1)."
)

set_para_text(750,
    "Performance (H4, partially supported): Median end-to-end latency was "
    "47ms, within the 50ms threshold; however, 95th-percentile latency "
    "reached 2.1 seconds, driven by external enrichment calls (GeoIP, "
    "WiGLE, SIEM) absent from the single-hop baselines (13-15ms median for "
    "ablation, Ahmadi, and Phani). This is a genuine, honest limitation "
    "rather than a uniformly satisfied hypothesis: typical-case latency "
    "supports H4, worst-case latency does not."
)

set_para_text(751,
    "Usability: The proposed framework's step-up rate was 70.79% on the "
    "evaluation set, reflecting its high true-positive rate against a "
    "STRIDE-injected, attack-heavy dataset (95% malicious by construction) "
    "rather than a before/after reduction — no baseline measurement of "
    "step-up rate without the validation layer exists for the proposed "
    "framework's own signals other than the ablation configuration, which "
    "showed both a lower step-up rate (42.60%) and a much higher, "
    "less-trustworthy FPR (20.95%). A session-continuity metric was not "
    "measured this cycle (see Limitations)."
)

set_para_text(752,
    "SIEM Integration (H3): SIEM correlation classifies live session "
    "anomalies into STRIDE categories with severity levels and is unique "
    "to the proposed framework among all evaluated configurations (Figure "
    "4.4). A controlled experiment isolating SIEM's specific TPR "
    "contribution was not performed this cycle."
)

set_para_text(753,
    "Privacy (H5): The proposed framework implements HMAC-SHA-256 hashing "
    "of contextual identifiers at ingestion and a bounded retention window, "
    "consistent with data-minimisation principles. A formal, independent "
    "privacy-leakage audit was not performed this cycle — this remains an "
    "implemented mechanism rather than a measured result."
)

set_para_text(762,
    "RQ1: Multi-source validation of contextual signals — cross-checking "
    "GPS, IP geolocation, and Wi-Fi BSSID for geographic consistency, and "
    "validating device/TLS signals — improved authentication accuracy "
    "across all metrics. The proposed framework achieved 88.30% TPR and "
    "2.86% FPR, compared to 21.86%/5.71% for Ahmadi (2025) and 6.11%/1.90% "
    "for Phani Kumar Kanuri (2025), a statistically significant "
    "improvement in both directions (McNemar's test, p < 0.001)."
)

set_para_text(763,
    "RQ2: Quality-weighted signal integration produced a false-positive "
    "rate of 2.86%, substantially lower than the ablation configuration's "
    "20.95% (validation layer disabled) and both re-implemented baselines "
    "(5.71% and 1.90%—though Phani's lower FPR reflects its extremely "
    "conservative decision rule, which also produces a 6.11% TPR; low FPR "
    "alongside low TPR is not evidence of better discrimination). The "
    "proposed framework achieves the best FPR/TPR balance of any "
    "evaluated configuration (F1 = 0.937)."
)

set_para_text(764,
    "RQ3: Real-time SIEM integration provides STRIDE-classified, "
    "severity-ranked alerting unavailable in any baseline configuration "
    "(Figure 4.4). Its specific quantitative contribution to detection "
    "accuracy under active threat conditions was not isolated in a "
    "controlled experiment this cycle."
)

set_para_text(765,
    "RQ4: The framework's median end-to-end latency (47ms) satisfies a "
    "50ms interactive-authentication threshold; 95th-percentile latency "
    "(2.1s) does not, driven by external enrichment calls absent from the "
    "single-hop re-implemented baselines (13-15ms median). This "
    "represents a measurable, architecturally expected cost of multi-"
    "source cross-validation that should be weighed against its "
    "associated accuracy gains rather than treated as uniformly bounded."
)

set_para_text(766,
    "RQ5: The framework balances security, usability, and privacy through "
    "three mechanisms: quality-weighted scoring that measurably reduces "
    "false-positive escalations relative to an unvalidated configuration; "
    "freshness-based graceful degradation intended to maintain "
    "functionality under signal uncertainty; and HMAC-SHA-256 hashing "
    "with data minimisation as an implemented (though not independently "
    "audited) privacy safeguard."
)

set_para_text(768,
    "The findings are subject to five constraints. First, the evaluation "
    "dataset (CIC-IDS2018, supplemented by RBA for real-world spoofing "
    "ground truth) is predominantly network/protocol-layer in its attack "
    "taxonomy, which limits how much any context-validation framework — "
    "proposed or baseline — can be expected to detect certain attack "
    "categories; this is a dataset-fit limitation rather than a framework "
    "weakness (Section 4.9). Second, cross-study comparison against the "
    "baseline papers' own self-reported figures (92.7% and 96.8% "
    "respectively) is not meaningful, since neither paper releases its "
    "evaluation dataset or methodology for independent verification — this "
    "thesis's comparison is instead a faithful re-implementation tested "
    "against a real, disclosed dataset, which is a different and stronger "
    "standard of evidence, not a directly comparable number. Third, "
    "endpoint telemetry was generated by simulation rather than real "
    "devices, and real-world signal distributions may differ. Fourth, the "
    "framework was evaluated by a single researcher without independent "
    "replication. Fifth, several results reported in earlier drafts of "
    "this chapter (granular component-level ablation, SIEM's specific TPR "
    "contribution, a formal privacy-leakage audit, and updated network-"
    "condition sensitivity figures) were not actually measured and have "
    "been removed or rescoped as future work rather than reported as "
    "findings."
)

set_para_text(777,
    "The perimeter model of enterprise security has been rendered obsolete "
    "by distributed, remote-first work environments. Zero Trust "
    "Architecture and adaptive MFA address this shift, but their "
    "effectiveness is contingent on the reliability of the contextual "
    "signals that underpin enforcement decisions. This study has shown "
    "that validating, quality-weighting, and integrating those signals "
    "before they influence authentication outcomes produces measurable, "
    "statistically significant improvements across every security accuracy "
    "metric: 88.3% TPR, 2.86% FPR, and F1 = 0.937, outperforming both "
    "re-implemented published baseline frameworks and an ablation "
    "configuration on the same dataset under identical conditions. The "
    "combination of multi-source cross-validation, real-time SIEM "
    "integration, and embedded privacy-preserving mechanisms constitutes a "
    "deployable, principled, evidence-based approach to context-aware Zero "
    "Trust authentication for the remote work era."
)

d.save(PATH)
print("Saved paragraph edits.")
